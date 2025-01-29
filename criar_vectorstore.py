from docx import Document
from langchain.docstore.document import Document as LangChainDocument
from langchain_chroma import Chroma
from langchain_ollama import OllamaEmbeddings
import os

# Diretório do VectorStore
persist_directory = "./chroma_db"

# Caminho fixo do arquivo .docx
file_path = r"C:\Users\isabella.silveira\Downloads\GEN_AI\books\Resumo_A_Quarta_Revolucao_Industrial.docx"

def create_documents(file_path):
    """
    Processa um arquivo .docx e retorna uma lista de objetos LangChainDocument.
    """
    try:
        doc = Document(file_path)
        documents = []
        for para in doc.paragraphs:
            if para.text.strip():  # Ignorar parágrafos vazios
                documents.append(LangChainDocument(page_content=para.text.strip()))
        return documents
    except Exception as e:
        print(f"Erro ao processar o arquivo: {e}")
        return []

def create_vectorstore(file_path):
    """
    Cria e salva o VectorStore no diretório persistente.
    """
    try:
        # Verificar se o arquivo existe
        if not os.path.exists(file_path):
            print(f"Erro: O arquivo '{file_path}' não foi encontrado.")
            return

        # Criar documentos a partir do arquivo
        documents = create_documents(file_path)
        if not documents:
            print("Nenhum documento válido encontrado no arquivo.")
            return

        # Criar e salvar o VectorStore
        vectorstore = Chroma.from_documents(
            documents=documents,
            embedding=OllamaEmbeddings(model="llama3"),
            persist_directory=persist_directory,
        )
        print("VectorStore criado e salvo com sucesso!")
    except Exception as e:
        print(f"Erro ao criar o VectorStore: {e}")

if __name__ == "__main__":
    create_vectorstore(file_path)
